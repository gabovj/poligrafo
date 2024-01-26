import streamlit as st
from pymongo import MongoClient
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Use your MongoDB Atlas connection string here
mongo_uri = st.secrets["mongo_uri"]
# Create a MongoClient using the provided URI
client = MongoClient(mongo_uri)
# Specify the database and collection
db = client["blanca"]
collection = db["poligrafia"]

def show_id():
    # MongoDB setup
    mongo_uri = st.secrets["mongo_uri"]

    # Define the fields to be displayed, excluding '_id'
    fields = {"_id": 0, "folio": 1, "name": 1, "timestamp": 1}

    # Date input for start and end date with None as default
    st.write("##### Selecciona un rango de fechas:")
    col1, col2 = st.columns(2)
    with col1:
        start_date = st.date_input("Fecha de inicio", value=None, key="start_date")
    with col2:
        end_date = st.date_input("Fecha final", value=None, key="end_date")

    # Proceed only if both dates are defined
    if start_date and end_date:
        # Ensure end_date is after start_date
        if start_date > end_date:
            st.error('Error: End date must fall after start date.')
            return
        
        # Connect to MongoDB using a context manager
        with MongoClient(mongo_uri) as client:
            db = client["blanca"]
            collection = db["poligrafia"]
            
            # Query to fetch data between the selected dates
            query = {
                "timestamp": {
                    "$gte": datetime.combine(start_date, datetime.min.time()),
                    "$lte": datetime.combine(end_date, datetime.max.time())
                }
            }
            
            # Fetch data from MongoDB with only the specified fields, filtered by date range
            data = list(collection.find(query, fields).sort("timestamp", -1))
            
            # Re-order the data to ensure the column order
            ordered_data = [
                {
                    "folio": item.get("folio", "<NA>"),
                    "name": item.get("name", "<NA>"),
                    "timestamp": item.get("timestamp", "<NA>")
                } for item in data
            ]

        # Display the data in a table if there is data to display
        if ordered_data:
            st.table(ordered_data)
        else:
            st.warning("No data available for the selected date range.")
    else:
        # Message to prompt the user to select both dates
        st.info("Please select both a start and an end date to display the data.")

def find_by_ID():
    col1, col2 = st.columns([1.3,.7])
    with col1:
        id_db = st.text_input(':orange[Ingresa el Folio]')
        
    with col2:
        st.markdown('####')
        find_by_id_button = st.button(':orange[Buscar información]')

    if find_by_id_button:
        existing_data = collection.find_one({"folio": id_db})

        if existing_data is not None:
            formatted_timestamp = existing_data['timestamp'].strftime("%d-%b-%Y").upper()
            st.write(f'Fecha de evaluación: :orange[{formatted_timestamp}]')
            
            with st.expander("#### :orange[Datos Generales]"):
                col1, col2 = st.columns([.7, .3])
                with col1:
                    if 'name' in existing_data:
                        st.write(f'##### Nombre: :orange[{existing_data["name"]}]')
                    else:
                        st.write('##### Nombre: :red[NO DATA]')

                with col2:
                    if 'dob' in existing_data:
                        dob = existing_data['dob']
                        dob_datetime = datetime.strptime(dob, "%Y-%m-%d")
                        now = datetime.now()
                        age = now.year - dob_datetime.year
                        if (now.month, now.day) < (dob_datetime.month, dob_datetime.day):
                            age -= 1
                        st.write(f'##### Edad: :orange[{age} años]')
                    else:
                        st.write('##### Edad: :red[NO DATA]')

                col1, col2, col3 = st.columns(3)
                with col1:
                    if 'dob' in existing_data:
                        formatted_dob = dob_datetime.strftime("%d-%b-%Y").upper()
                        st.write(f'Nacimiento: :orange[{formatted_dob}]')
                    else:
                        st.write('Nacimiento: :red[NO DATA]')

                with col2:
                    if 'marriage_status' in existing_data:
                        st.write(f'Estado civil: :orange[{existing_data["marriage_status"]}]')
                    else:
                        st.write('Estado civil: :red[NO DATA]')

                with col3:
                    if 'birth_place' in existing_data:
                        st.write(f'Lugar nacimiento: :orange[{existing_data["birth_place"]}]')
                    else:
                        st.write('Lugar nacimiento: :red[NO DATA]')
                
                col1, col2, col3 = st.columns([.5, .2, .3])
                with col1:
                    if 'street_adress' in existing_data:
                        st.write(f'Calle: :orange[{existing_data["street_adress"]}]')
                    else:
                        st.write('Calle: :red[NO DATA]')

                with col2:
                    if 'street_number' in existing_data:
                        st.write(f'Número: :orange[{existing_data["street_number"]}]')
                    else:
                        st.write('Número: :red[NO DATA]')

                with col3:
                    if 'neighborhood' in existing_data:
                        st.write(f'Colonia: :orange[{existing_data["neighborhood"]}]')
                    else:
                        st.write('Colonia: :red[NO DATA]')

                col1, col2 = st.columns([.4, .6])
                with col1:
                    if 'phone' in existing_data:
                        st.write(f'Teléfono: :orange[{existing_data["phone"]}]')
                    else:
                        st.write('Teléfono: :red[NO DATA]')

                    if 'academic_grade' in existing_data:
                        st.write(f'Grado Académico: :orange[{existing_data["academic_grade"]}]')
                    else:
                        st.write('Grado Académico: :red[NO DATA]')

                    if 'couple_name' in existing_data:
                        st.write(f'Nombre de su pareja: :orange[{existing_data["couple_name"]}]')
                    else:
                        st.write('Nombre de su pareja: :red[NO DATA]')

                with col2:
                    if 'email' in existing_data:
                        st.write(f'Email: :orange[{existing_data["email"]}]')
                    else:
                        st.write('Email: :red[NO DATA]')

                    if 'couple_ocupation' in existing_data:
                        st.write(f'Ocupación de su pareja: :orange[{existing_data["couple_ocupation"]}]')
                    else:
                        st.write('Ocupación de su pareja: :red[NO DATA]')

                    if 'children_number' in existing_data:
                        st.write(f'Núm de hijos: :orange[{existing_data["children_number"]}]')
                    else:
                        st.write('Núm de hijos: :red[NO DATA]')

                if 'children_names' in existing_data:
                    st.write(f'Nombre de hijos: :orange[{existing_data["children_names"]}]')
                else:
                    st.write('Nombre de hijos: :red[NO DATA]')

                if 'economic_dependents' in existing_data:
                    st.write(f'Dependientes económicos: :orange[{existing_data["economic_dependents"]}]')
                else:
                    st.write('Dependientes económicos: :red[NO DATA]')

                col1, col2 = st.columns(2)
                with col1:
                    if 'polygraph_experiences' in existing_data:
                        st.write(f'Experiencias Poligráficas: :orange[{existing_data["polygraph_experiences"]}]')
                    else:
                        st.write('Experiencias Poligráficas: :red[NO DATA]')

                with col2:
                    if 'polygraph_experiences_results' in existing_data:
                        st.write(f'Resultados: :orange[{existing_data["polygraph_experiences_results"]}]')
                    else:
                        st.write('Resultados: :red[NO DATA]')
           
            if 'labor_data' in existing_data:
                with st.expander("#### :orange[Datos Laborales]"):
                    for labor in existing_data['labor_data']:
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write(f'Empresa: :orange[{labor["company_name"]}]')
                        with col2:
                            st.write(f'Puesto: :orange[{labor["position"]}]')
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            start_date = labor['start_date']
                            start_date_str = datetime.strptime(start_date, "%Y-%m-%d")
                            formatted_start_date = start_date_str.strftime("%d-%b-%Y").upper()
                            st.write(f'Ingreso: :orange[{formatted_start_date}]')
                        with col2:
                            end_date = labor['end_date']
                            end_date_str = datetime.strptime(end_date, "%Y-%m-%d")
                            formatted_end_date = end_date_str.strftime("%d-%b-%Y").upper()
                            st.write(f'Salida: :orange[{formatted_end_date}]')
                        with col3:
                            st.write(f'Salario Mensual: :orange[{labor["monthly_salary"]}]')
                        st.write(f'Motivo de Salida: :orange[{labor["reason_exit"]}]')
                        st.write(f'Mala experiencia: :orange[{labor["worst_job_experience"]}]')
                        job_duties_with_breaks = labor["job_duties"].replace("\n", "<br>")
                        st.markdown(f'Actividades: :orange[{job_duties_with_breaks}]', unsafe_allow_html=True)
                        st.divider()
            else:
                with st.expander("#### :orange[Datos Laborales]"):
                    st.warning("No hay datos laborales disponibles.")

            with st.expander("#### :orange[Historia Médica]"):
                col1, col2 = st.columns(2)
                with col1:
                    medical_fields = [
                        ("health_condition", "Condición General de Salud"),
                        ("current_physical_discomfort", "Malestar físico actual"),
                        ("chronic_disease", "Enfermedad crónica"),
                        ("last_meal_time", "Última comida"),
                        ("nervous_treatment", "Tratamiento Nervioso"),
                        ("psychologist_treatment", "Tratamiento Psicológico o Psiquiátrico"),
                        ("recent_hospitalizations", "Hospitalizaciones recientes"),
                        ("recent_surgeries", "Cirugías recientes")
                    ]
                    for field, description in medical_fields:
                        if field in existing_data:
                            st.write(f'{description}: :orange[{existing_data[field]}]')
                        else:
                            st.write(f'{description}: :red[NO DATA]')
                
                with col2:
                    specify_fields = [
                        ("specify_health_condition", "Especificar Condición General de Salud"),
                        ("specify_current_physical_discomfort", "Especificar Malestar físico actual"),
                        ("recent_medical_treatment", "Tratamiento médico reciente"),
                        ("sleep_hours", "Horas de sueño"),
                        ("specify_nervous_treatment", "Especificar Tratamiento Nervioso"),
                        ("specify_psychologist_treatment", "Especificar Tratamiento Psicológico o Psiquiátrico"),
                        ("specify_recent_hospitalizations", "Especificar Hospitalizaciones recientes"),
                        ("specify_recent_surgeries", "Especificar Cirugías recientes")
                    ]
                    for field, description in specify_fields:
                        if field in existing_data:
                            st.write(f'{description}: :orange[{existing_data[field]}]')
                        else:
                            st.write(f'{description}: :red[NO DATA]')

                    # Continuing from the specify_fields loop
                boolean_fields_with_specification = [
                    ("respiratory", "Respiratorio", "specify_respiratory"),
                    ("heart", "Corazón", "specify_heart"),
                    ("blood_pressure", "Presión", "specify_blood_pressure"),
                    ("diabetes", "Diabetes", "specify_diabetes"),
                    ("fainting", "Desmayos", "specify_fainting"),
                    ("convulsion", "Convulsiones", "specify_convulsion"),
                    ("allergies", "Alergias", "specify_allergies"),
                    ("digestives", "Digestivos", "specify_digestives"),
                    ("pregnancy", "Embarazo", "specify_pregnancy"),
                    ("migraine", "Migraña", "specify_migraine"),
                    ("injuries", "Heridas", "specify_injuries"),
                    ("others", "Otros", "specify_others")
                ]

                for field, description, specify_field in boolean_fields_with_specification:
                    col1, col2 = st.columns([.25, .75])
                    with col1:
                        if field in existing_data:
                            st.write(f'{description}: {boolean_to_yes_no(existing_data[field])}')
                        else:
                            st.write(f'{description}: :red[NO DATA]')
                    
                    with col2:
                        if specify_field in existing_data:
                            st.write(f'Especificar {description}: :orange[{existing_data[specify_field]}]')
                        else:
                            st.write(f'Especificar {description}: :red[NO DATA]')
     
            with st.expander("#### :orange[Aspectos Emocionales]"):
                col1, col2 = st.columns(2)
                with col1:
                    if 'father_name' in existing_data:
                        st.write(f'Nombre del Padre: :orange[{existing_data["father_name"]}]')
                    else:
                        st.write(f'Nombre del Padre: :red[NO DATA]')
                    if 'mother_name' in existing_data:
                        st.write(f'Nombre de la Madre: :orange[{existing_data["mother_name"]}]')
                    else:
                        st.write(f'Nombre de la Madre: :red[NO DATA]')
                    
                with col2:
                    if "father_ocupation" in existing_data:
                        st.write(f'Ocupación: :orange[{existing_data["father_ocupation"]}]')
                    else:
                        st.write(f'Ocupación: :red[NO DATA]')
                    if "mother_ocupation" in existing_data:
                        st.write(f'Ocupación: :orange[{existing_data["mother_ocupation"]}]')
                    else:
                        st.write(f'Ocupación: :red[NO DATA]')
                col1, col2 = st.columns([.3,.7])
                with col1:
                    if "num_brothers" in existing_data:
                        st.write(f'Num de hermanos: :orange[{existing_data["num_brothers"]}]')
                    else: 
                        st.write(f'Num de hermanos: :red[NO DATA]')
                with col2:
                    if "brothers_name" in existing_data:
                        name_brothers_with_breaks = existing_data["brothers_name"].replace("\n", "<br>")
                        st.markdown(f'Nombre hermanos: :orange[{name_brothers_with_breaks}]', unsafe_allow_html=True)
                    else:
                        st.markdown(f'Nombre hermanos: :red[NO DATA]')
                if "brothers_ocupation" in existing_data:
                    brothers_occupation_with_breaks = existing_data["brothers_ocupation"].replace("\n", "<br>")
                    st.markdown(f'Ocupación hermanos: :orange[{brothers_occupation_with_breaks}]', unsafe_allow_html=True)
                else:
                    st.markdown(f'Ocupación hermanos: :red[NO DATA]')
                col1, col2 = st.columns([.4,.6])
                with col1:
                    if "who_do_you_live_with" in existing_data:
                        st.write(f'Con quién vives: :orange[{existing_data["who_do_you_live_with"]}]')
                    else:
                        st.write(f'Con quién vives: :red[NO DATA]')
                with col2:
                    if "best_life_experience" in existing_data:
                        st.write(f'Mejor experiencia: :orange[{existing_data["best_life_experience"]}]')
                    else:
                        st.write(f'Mejor experiencia: :red[NO DATA]')
                if "worst_life_experience" in existing_data:
                    st.write(f'Experiencia desagradable: :orange[{existing_data["worst_life_experience"]}]')
                else:
                    st.write(f'Experiencia desagradable: :red[NO DATA]')
            
            with st.expander("#### :orange[Aspectos Económicos]"):
                col1, col2 = st.columns(2)

                with col1:
                    if 'debts' in existing_data:
                        st.write(f'Deudas económicas: :orange[{existing_data["debts"]}]')
                    else:
                        st.write('Deudas económicas: :red[NO DATA]')

                    if 'banks' in existing_data:
                        st.write(f'Bancos: :orange[{existing_data["banks"]}]')
                    else:
                        st.write('Bancos: :red[NO DATA]')

                    if 'tarjetas_departamentales' in existing_data:
                        st.write(f'Tarjetas departamentales: :orange[{existing_data["tarjetas_departamentales"]}]')
                    else:
                        st.write('Tarjetas departamentales: :red[NO DATA]')

                    if 'empenos' in existing_data:
                        st.write(f'Empeños: :orange[{existing_data["empenos"]}]')
                    else:
                        st.write('Empeños: :red[NO DATA]')

                with col2:
                    if 'specify_debts' in existing_data:
                        st.write(f'Especificar: :orange[{existing_data["specify_debts"]}]')
                    else:
                        st.write('Especificar: :red[NO DATA]')

                    if 'bank_debt' in existing_data:
                        st.write(f'Deuda Bancaria: :orange[{existing_data["bank_debt"]}]')
                    else:
                        st.write('Deuda Bancaria: :red[NO DATA]')

                    if 'tarjetas_departamentales_debt' in existing_data:
                        st.write(f'Deuda Tarjetas Departamentales: :orange[{existing_data["tarjetas_departamentales_debt"]}]')
                    else:
                        st.write('Deuda Tarjetas Departamentales: :red[NO DATA]')

                    if 'empenos_debt' in existing_data:
                        st.write(f'Cuantía de Empeños: :orange[{existing_data["empenos_debt"]}]')
                    else:
                        st.write('Cuantía de Empeños: :red[NO DATA]')

                st.write('##### Bienes Patrimoniales')
                col1, col2, col3 = st.columns(3)

                with col1:
                    if 'cars' in existing_data:
                        st.write(f'Autos: :orange[{existing_data["cars"]}]')
                    else:
                        st.write('Autos: :red[NO DATA]')

                    if 'house' in existing_data:
                        st.write(f'Casa: :orange[{existing_data["house"]}]')
                    else:
                        st.write('Casa: :red[NO DATA]')

                with col2:
                    if 'properties' in existing_data:
                        st.write(f'Inmueble: :orange[{existing_data["properties"]}]')
                    else:
                        st.write('Inmueble: :red[NO DATA]')

                    if 'investments' in existing_data:
                        st.write(f'Inversiones: :orange[{existing_data["investments"]}]')
                    else:
                        st.write('Inversiones: :red[NO DATA]')

                with col3:
                    if 'business' in existing_data:
                        st.write(f'Negocio: :orange[{existing_data["business"]}]')
                    else:
                        st.write('Negocio: :red[NO DATA]')

            with st.expander("#### :orange[Información Laboral]"):
                laboral_questions = [
                    "¿Se te ha pedido renuncia de cualquiera de tus empleos anteriores?",
                    "¿Se te ha acusado alguna vez de deshonestidad de cualquiera de tus empleos anteriores?",
                    "¿Has violado los reglamentos establecidos de cualquier empleo?",
                    "¿Has usado tu puesto o tus funciones para obtener beneficios en forma ilícita o no autorizada?",
                    "¿Has hecho uso indebido de información confidencial de cualquiera de tus trabajos?",
                    "¿Cuál ha sido el problema laboral más serio que has tenido?",
                    "¿Has tenido discusiones serias con superiores?",
                    "¿Cuáles han sido las faltas administrativas más serias que has cometido de empleos anteriores?",
                    "¿Se ha levantado algún acta administrativa o algún proceso laboral?",
                    "¿Has aceptado sobornos de cualquier tipo de cualquiera de tus empleos anteriores?",
                    "¿Has saboteado tu trabajo o el de tus compañeros?",
                    "¿Se ha levantado alguna demanda laboral en tu contra?",
                    "¿Has demandado laboralmente a la empresa o alguno de tus jefes?",
                    "¿Supiste de personas de tus trabajos anteriores que hayan realizando actos ilegales?",
                    "¿Has participado de algún robo en alguna empresa?",
                    "¿Te has visto comprometido en realizar alguna conducta que afecte los intereses de la empresa?",
                    "¿Has falsificado o alterado de manera intencional documentos de tu empresa?",
                    "¿Has hecho uso indebido de los recursos de la empresa?"
                ]

                for i, question in enumerate(laboral_questions, start=1):
                    key = f'info_laboral_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')

            with st.expander("#### :orange[Actividades Delictivas]"):
                delictivas_questions = [
                    "¿Conoces a delincuentes o grupos delictivos de cualquier tipo?",
                    "¿Has sido detenido o cuestionado por la policía por cualquier motivo?",
                    "¿Has estado detenido en delegaciones, reclusorios, cárceles o penitenciarias?",
                    "¿Has hecho algo por lo que pudiste haber sido detenido?",
                    "¿Alguien de tu familia ha estado detenido en reclusorios, cárceles o penitenciarias?",
                    "¿Has tenido cargos por cualquier falta cometida?",
                    "¿Utilizas armas de fuego?",
                    "¿Venta, comercialización de armas de fuego?",
                    "¿Le has causado lesiones a personas con armas de fuego?",
                    "¿Tienes antecedentes penales?",
                    "¿Has ayudado a alguien a cometer alguna actividad ilícita, delito o acto vandálico?",
                    "¿Se ha levantado alguna demanda civil, mercantil o judicial, en tu contra?",
                    "¿Has cometido fraude, extorsión o chantaje?",
                    "¿Has estado involucrado en algún secuestro?"
                ]

                for i, question in enumerate(delictivas_questions, start=1):
                    key = f'actividades_delictivas_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')

            with st.expander("#### :orange[Hábitos Personales]"):
                habitos_personales_questions = [
                    "¿Cuándo fue la última vez que bebiste?",
                    "¿Cuánto bebes al mes en promedio?",
                    "¿Cuándo fue la última vez que olvidaste lo ocurrido mientras bebías?",
                    "¿Cuándo fue la última vez que manejaste en estado de ebriedad?",
                    "¿Ha sido detenido por conducir en estado de embriaguez vehículos de la empresa?",
                    "¿Te has presentado a laborar en estado de ebriedad, con aliento alcohólico o con resaca?",
                    "¿Has faltado a tu trabajo por haber ingerido alcohol?",
                    "¿Has ingerido bebidas alcohólicas en horas de trabajo?"
                ]

                for i, question in enumerate(habitos_personales_questions, start=1):
                    key = f'habitos_personales_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')

            with st.expander("#### :orange[Drogas]"):
                drogas_questions = [
                    "¿Consumes alguna droga ilegal?",
                    "¿Has probado o experimentado con cualquier droga ilegal?",
                    "¿Has tenido algún tipo de contacto con droga ilegal en su trabajo?",
                    "¿Has comprado o vendido cualquier droga ilegal?",
                    "¿Has transportado o distribuido alguna droga ilegal?",
                    "¿Tienes contactos con personas que vendan o distribuyan droga?",
                    "¿Sabe de alguna persona de la empresa que tenga relación con drogas ilegales?"
                ]

                for i, question in enumerate(drogas_questions, start=1):
                    key = f'drogas_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')

            with st.expander("#### :orange[Antecedentes de Tránsito]"):
                tansito_questions =['¿Tienes antecedentes de tránsito?']
                for i, question in enumerate(tansito_questions, start=1):
                    key = f'transito_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')

            with st.expander("#### :orange[Conclusión o Cierre]"):
                conclusion_questions = [
                    "¿Has tenido alguna experiencia en tu trabajo por la que pudieras ser chantajeado posteriormente?",
                    "¿Pensaste que el día de hoy iba a hacerte alguna otra pregunta de algo que no hayas comentado hasta el momento?",
                    "¿Falseaste u omitiste deliberadamente alguna información que te haya solicitado el día de hoy?",
                    "¿Hay algo en tu mente en este momento que te preocupe para no poder tener un buen examen?",
                    "¿Hay algo que quieras decirme en este momento y estés dudando si hacerlo o no?",
                    "Observaciones"
                ]

                for i, question in enumerate(conclusion_questions, start=1):
                    key = f'conclusion_{i}'
                    if key in existing_data:
                        st.write(f'{question} :orange[{existing_data[key]}]')
                    else:
                        st.write(f'{question} :red[NO DATA]')
            
        else:
            st.error('No se encontraron datos para el folio proporcionado.')
        
        

def boolean_to_yes_no(value):
    return ':white_check_mark:' if value else ':x:'
    
def create_paragraph(doc, text_parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment

    for text, underline in text_parts:
        run = paragraph.add_run(text)
        if underline:
            run.underline = True

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

    return paragraph

def create_docx(variables):
    print(variables)
    # Mapping of month numbers to Spanish month names
    spanish_months = {1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril', 5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto', 9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'}

    # Create a new Document
    doc = Document()
    # Set the default font to Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Add a logo to the header
    try:
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_paragraph.add_run().add_picture("test_logo.png", width=Inches(1.0))
    except FileNotFoundError:
        print("Logo file not found.")

    # Add and center the heading
    heading = doc.add_heading("AUTORIZACIÓN Y LIBERACIÓN DEL EXÁMEN DE POLÍGRAFO", level=1)
    heading.style = doc.styles['Heading 1']
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Modify the heading font
    heading_font = heading.style.font
    heading_font.name = 'Arial'  # Ensure heading also uses Arial
    heading_font.size = Pt(13)   # Set font size to 13 points

    # Date paragraph
    datetime_object = variables.get('timestamp', None)
    if datetime_object:
        date_text = f'Puebla, Puebla a {datetime_object.day} de {spanish_months[datetime_object.month]} de {datetime_object.year}'
        create_paragraph(doc, [(date_text, True)], WD_ALIGN_PARAGRAPH.RIGHT)

     # Adding paragraphs with specific underlined parts
    name = variables.get('name', 'N/A')
    id_number = variables.get('id_number','N/A')  
    hiring_company = variables.get('hiring_company', 'N/A')  
    key_poligraphist = variables.get('key_poligraphist', 'N/A')  
    texts = [
        [("Yo ", False), (name, True), (" con número de identificación: ", False),(id_number, True)],
        [("Por la presente autorizo al profesional, quien es examinador de polígrafo calificado, para administrarme un examen de polígrafo en la fecha indicada. Además, entiendo que puedo cancelar los procedimientos de evaluación en cualquier momento, y que será detenido el examen o entrevista si lo deseo.", False)],
        [("Estoy de acuerdo en tomar este examen, sin ninguna promesa de recompensa de cualquier tipo; incluida la alteración u omisión del resultado o información producto de la evaluación.", False)],
        [("Entiendo que con base a cómo yo genere los datos de la prueba del polígrafo, la opinión experta del examinador puede ser: que he sido del todo veraz o no; y que he cooperado o no completamente con los procedimientos de la prueba.", False)],
        [("Tengo salud física y mental adecuada (estable) para completar esta examinación.", False)],
        [("Entiendo que el equipo de examinación y los principios fisiológicos de la prueba, se explicarán a mi satisfacción antes de comenzar el examen.", False)],
        [("Entiendo que todas las preguntas que se me harán durante el examen serán leídas y revisadas conmigo antes de comenzar la prueba, y que voy a tener la oportunidad de pedir una aclaración antes de empezar con el mismo.", False)],
        [("Registro de examinación: Audio - Video", False)],
        [("Autorizo la entrega de resultados y la información de este examen, de manera: oral, escrita y electrónica; incluyendo el contenido del video y audio grabado de entrevista; los resultados, informe sumario y opiniones a: ", False), (hiring_company, True)],
        [("Yo, en pleno uso de mis facultades, entiendo que esta prueba la realizo por mi propia voluntad ya que autoricé todos los procedimientos por escrito, se me explicó que los instrumentos van a ser colocados sobre mi cuerpo y que ninguno de estos produce daño. Por lo tanto, eximo de toda responsabilidad civil y penal en contra de; el poligrafista con clave: ", False), (key_poligraphist, True), ("; la empresa a la que pertenece éste y la empresa solicitante, de mi evaluación ya que no me han obligado a presentar la prueba poligráfica además de tener conocimiento que en cualquier momento puedo abandonar este lugar y con ello no se violentan mis garantías individuales ni mis derechos constitucionales.", False)]
    ]

    for text_parts in texts:
        create_paragraph(doc, text_parts)

    # Add 3 empty lines
    for _ in range(3):
        doc.add_paragraph()

    # Add a table for signatures
    signature_table = doc.add_table(rows=1, cols=2)
    signature_table.autofit = False

    # Set column widths - Adjust these values as needed
    signature_table.columns[0].width = Inches(3.25)
    signature_table.columns[1].width = Inches(3.25)

    # Add underlined spaces for signatures in each cell
    signature_line = '_ ' * 15  # Adjust as necessary for signature line length
    cell_left = signature_table.cell(0, 0)
    cell_right = signature_table.cell(0, 1)

    cell_left_paragraph = cell_left.paragraphs[0]
    cell_right_paragraph = cell_right.paragraphs[0]

    # Add signature line and text for the left cell
    cell_left_paragraph.add_run(signature_line).underline = True
    cell_left_paragraph.add_run("\nFIRMA DEL EVALUADO")  # New line and text

    # Add signature line and text for the right cell
    cell_right_paragraph.add_run(signature_line).underline = True
    cell_right_paragraph.add_run("\nFIRMA DEL POLIGRAFISTA")  # New line and text

    # Align the text in the cells
    cell_left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_right_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add 2 empty lines
    for _ in range(2):
        doc.add_paragraph()

    # Add the current time above the footer
    current_time = datetime.now().strftime("%H:%M")
    time_paragraph = doc.add_paragraph(f"Hora de inicio: {current_time}")
    time_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Footer
    footer_text = "INFORMACIÓN CONFIDENCIAL. El presente documento fue elaborado a petición del cliente, realizado mediante los estándares establecidos por la Asociación Mexicana de Poligrafía, el proceso fue video grabado para fines de control de calidad, utilizo un instrumento digital de poligrafía para la información. Quien ha aceptado que toda información contenida en el mismo es confidencial y de su interés exclusivo y privado. Cualquier uso distinto al descrito, ya sea comunicación verbal, publicación o reproducción total o parcial sin la aprobación escrita del Centro Especializado en Poligrafía e investigación Estratégica queda estrictamente prohibido. A sí mismo el cliente acepta de conformidad la responsabilidad imputable ocasionada por el uso indebido de este documento."
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    footer_paragraph.add_run(footer_text).font.size = Pt(5)

    # Save the document to a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream

